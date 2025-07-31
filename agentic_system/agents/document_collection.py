"""
Document Collection Agent for the Agentic System.

This agent handles document parsing, processing, and indexing using:
- PDF processing via pdfplumber and PyMuPDF
- DOCX processing via python-docx
- LangChain integration with Chroma vector store
- Multi-modal content extraction (text, tables, images)
"""

import asyncio
import hashlib
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime
import json

# Document processing libraries
import pdfplumber
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_openai import AzureOpenAIEmbeddings
from langchain_community.vectorstores import Chroma

# Local imports
from ..core.models import (
    AgentMessage, DocumentMetadata, DocumentType, AgentType, 
    MessageType, AuditEntry, ProcessingStats
)
from ..core.agentic_config import config, logger


class DocumentProcessingError(Exception):
    """Custom exception for document processing errors."""
    pass


class DocumentCollectionAgent:
    """
    Agent responsible for document collection, parsing, and indexing.
    
    Capabilities:
    - Multi-format document parsing (PDF, DOCX, TXT)
    - Text chunking and embedding generation
    - Vector store indexing with Chroma
    - Multi-modal content extraction
    - Metadata management and audit trail
    """
    
    def __init__(self, agent_config: Optional[Dict[str, Any]] = None):
        """Initialize the Document Collection Agent."""
        self.agent_type = AgentType.DOCUMENT_COLLECTION
        self.config = config
        self.logger = logger.bind(agent=self.agent_type.value)
        
        # Initialize Azure OpenAI embeddings
        self.embeddings = AzureOpenAIEmbeddings(
            azure_endpoint=self.config.azure_openai.endpoint,
            api_key=self.config.azure_openai.api_key,
            api_version=self.config.azure_openai.api_version,
            azure_deployment=self.config.azure_openai.embedding_model,
        )
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.document_processing.chunk_size,
            chunk_overlap=self.config.document_processing.chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Initialize vector store
        self.vector_store = self._initialize_vector_store()
        
        # Processing statistics
        self.stats = ProcessingStats()
        
        self.logger.info("Document Collection Agent initialized")
    
    def _initialize_vector_store(self) -> Chroma:
        """Initialize or load existing Chroma vector store."""
        try:
            vector_store = Chroma(
                collection_name=self.config.vector_store.collection_name,
                embedding_function=self.embeddings,
                persist_directory=self.config.vector_store.persist_directory
            )
            self.logger.info("Vector store initialized", 
                           collection=self.config.vector_store.collection_name)
            return vector_store
        except Exception as e:
            self.logger.error("Failed to initialize vector store", error=str(e))
            raise DocumentProcessingError(f"Vector store initialization failed: {e}")
    
    async def process_document(self, file_path: Path) -> DocumentMetadata:
        """
        Process a single document and return metadata.
        
        Args:
            file_path: Path to the document to process
            
        Returns:
            DocumentMetadata: Metadata about the processed document
            
        Raises:
            DocumentProcessingError: If processing fails
        """
        start_time = datetime.now()
        
        try:
            # Validate file exists and format is supported
            if not file_path.exists():
                raise DocumentProcessingError(f"File not found: {file_path}")
            
            file_extension = file_path.suffix.lower()
            if file_extension not in self.config.document_processing.supported_formats:
                raise DocumentProcessingError(f"Unsupported file format: {file_extension}")
            
            # Create document metadata
            metadata = DocumentMetadata(
                file_path=file_path,
                document_type=DocumentType(file_extension[1:]),  # Remove dot
                file_size=file_path.stat().st_size,
                created_at=datetime.fromtimestamp(file_path.stat().st_ctime)
            )
            
            # Extract content based on file type
            if file_extension == ".pdf":
                content_data = await self._process_pdf(file_path)
            elif file_extension == ".docx":
                content_data = await self._process_docx(file_path)
            elif file_extension == ".txt":
                content_data = await self._process_txt(file_path)
            else:
                raise DocumentProcessingError(f"Handler not implemented for {file_extension}")
            
            # Update metadata with extraction results
            metadata.text_length = len(content_data["text"])
            metadata.has_tables = content_data.get("has_tables", False)
            metadata.has_images = content_data.get("has_images", False)
            metadata.page_count = content_data.get("page_count")
            metadata.content_hash = self._calculate_content_hash(content_data["text"])
            
            # Create chunks and add to vector store
            await self._add_to_vector_store(content_data["text"], metadata)
            
            # Mark as successfully processed
            metadata.processed_at = datetime.now()
            metadata.extraction_successful = True
            
            # Update statistics
            self.stats.total_documents += 1
            self.stats.processed_documents += 1
            processing_time = (datetime.now() - start_time).total_seconds()
            self.stats.average_processing_time = (
                (self.stats.average_processing_time * (self.stats.processed_documents - 1) + processing_time) 
                / self.stats.processed_documents
            )
            
            # Log audit entry
            await self._log_audit_entry("document_processed", {
                "file_path": str(file_path),
                "processing_time": processing_time,
                "success": True
            })
            
            self.logger.info("Document processed successfully", 
                           file_path=str(file_path), 
                           processing_time=processing_time)
            
            return metadata
            
        except Exception as e:
            # Update failure statistics
            self.stats.total_documents += 1
            self.stats.failed_documents += 1
            
            # Log audit entry for failure
            await self._log_audit_entry("document_processing_failed", {
                "file_path": str(file_path),
                "error": str(e),
                "success": False
            })
            
            self.logger.error("Document processing failed", 
                            file_path=str(file_path), 
                            error=str(e))
            raise DocumentProcessingError(f"Processing failed for {file_path}: {e}")
    
    async def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Process PDF document and extract text, tables, and images."""
        content_data = {
            "text": "",
            "has_tables": False,
            "has_images": False,
            "page_count": 0
        }
        
        try:
            # Use pdfplumber for text and table extraction
            with pdfplumber.open(file_path) as pdf:
                content_data["page_count"] = len(pdf.pages)
                all_text = []
                
                for page in pdf.pages:
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        all_text.append(page_text)
                    
                    # Check for tables
                    tables = page.extract_tables()
                    if tables:
                        content_data["has_tables"] = True
                        # Convert tables to text format
                        for table in tables:
                            table_text = "\n".join(["\t".join(row) for row in table if row])
                            all_text.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")
                
                content_data["text"] = "\n\n".join(all_text)
            
            # Use PyMuPDF for image extraction
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                image_list = page.get_images()
                if image_list:
                    content_data["has_images"] = True
                    break  # We just need to know if images exist
            
            doc.close()
            
        except Exception as e:
            self.logger.error("PDF processing failed", file_path=str(file_path), error=str(e))
            raise DocumentProcessingError(f"PDF processing failed: {e}")
        
        return content_data
    
    async def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Process DOCX document and extract text and tables."""
        content_data = {
            "text": "",
            "has_tables": False,
            "has_images": False,
            "page_count": None  # Not easily available for DOCX
        }
        
        try:
            doc = DocxDocument(file_path)
            all_text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    all_text.append(paragraph.text)
            
            # Extract tables
            if doc.tables:
                content_data["has_tables"] = True
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = "\t".join([cell.text for cell in row.cells])
                        table_text.append(row_text)
                    all_text.append(f"\n[TABLE]\n{chr(10).join(table_text)}\n[/TABLE]\n")
            
            # Check for images (inline shapes)
            if doc.inline_shapes:
                content_data["has_images"] = True
            
            content_data["text"] = "\n\n".join(all_text)
            
        except Exception as e:
            self.logger.error("DOCX processing failed", file_path=str(file_path), error=str(e))
            raise DocumentProcessingError(f"DOCX processing failed: {e}")
        
        return content_data
    
    async def _process_txt(self, file_path: Path) -> Dict[str, Any]:
        """Process plain text document."""
        content_data = {
            "text": "",
            "has_tables": False,
            "has_images": False,
            "page_count": None
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content_data["text"] = f.read()
        except Exception as e:
            self.logger.error("TXT processing failed", file_path=str(file_path), error=str(e))
            raise DocumentProcessingError(f"TXT processing failed: {e}")
        
        return content_data
    
    async def _add_to_vector_store(self, text: str, metadata: DocumentMetadata) -> None:
        """Add document content to vector store."""
        try:
            # Create text chunks
            chunks = self.text_splitter.split_text(text)
            
            # Create Document objects with metadata
            documents = []
            for i, chunk in enumerate(chunks):
                doc_metadata = {
                    "source": str(metadata.file_path),
                    "document_type": metadata.document_type.value,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "file_size": metadata.file_size,
                    "has_tables": metadata.has_tables,
                    "has_images": metadata.has_images,
                    "processed_at": metadata.processed_at.isoformat() if metadata.processed_at else None
                }
                
                documents.append(Document(
                    page_content=chunk,
                    metadata=doc_metadata
                ))
            
            # Add to vector store
            await asyncio.get_event_loop().run_in_executor(
                None, 
                lambda: self.vector_store.add_documents(documents)
            )
            
            self.logger.info("Document added to vector store", 
                           file_path=str(metadata.file_path),
                           chunk_count=len(chunks))
            
        except Exception as e:
            self.logger.error("Failed to add document to vector store", 
                            file_path=str(metadata.file_path), 
                            error=str(e))
            raise DocumentProcessingError(f"Vector store addition failed: {e}")
    
    def _calculate_content_hash(self, content: str) -> str:
        """Calculate SHA-256 hash of content for change detection."""
        return hashlib.sha256(content.encode('utf-8')).hexdigest()
    
    async def search_documents(self, query: str, top_k: Optional[int] = None) -> List[Document]:
        """
        Search documents in the vector store.
        
        Args:
            query: Search query
            top_k: Number of results to return (defaults to config value)
            
        Returns:
            List of relevant document chunks
        """
        if top_k is None:
            top_k = self.config.vector_store.top_k
        
        try:
            # Perform similarity search
            results = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: self.vector_store.similarity_search(
                    query, 
                    k=top_k
                )
            )
            
            self.logger.info("Document search completed", 
                           query=query, 
                           results_count=len(results))
            
            return results
            
        except Exception as e:
            self.logger.error("Document search failed", query=query, error=str(e))
            raise DocumentProcessingError(f"Search failed: {e}")
    
    async def process_document_batch(self, file_paths: List[Path]) -> List[DocumentMetadata]:
        """
        Process multiple documents concurrently.
        
        Args:
            file_paths: List of document paths to process
            
        Returns:
            List of document metadata for successfully processed documents
        """
        self.logger.info("Starting batch document processing", 
                        document_count=len(file_paths))
        
        # Process documents concurrently
        tasks = [self.process_document(path) for path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Separate successful results from errors
        successful_metadata = []
        for i, result in enumerate(results):
            if isinstance(result, DocumentMetadata):
                successful_metadata.append(result)
            else:
                self.logger.error("Batch processing error", 
                                file_path=str(file_paths[i]), 
                                error=str(result))
        
        self.logger.info("Batch processing completed", 
                        total_documents=len(file_paths),
                        successful=len(successful_metadata),
                        failed=len(file_paths) - len(successful_metadata))
        
        return successful_metadata
    
    async def _log_audit_entry(self, action: str, details: Dict[str, Any]) -> None:
        """Log audit entry for document processing actions."""
        if not self.config.agent_coordination.audit_enabled:
            return
        
        audit_entry = AuditEntry(
            entry_id=f"doc_{datetime.now().timestamp()}",
            agent_type=self.agent_type,
            action=action,
            details=details
        )
        
        # Write to audit log file
        try:
            with open(self.config.agent_coordination.audit_log_path, 'a') as f:
                f.write(audit_entry.to_json() + "\n")
        except Exception as e:
            self.logger.error("Failed to write audit entry", error=str(e))
    
    def get_processing_stats(self) -> ProcessingStats:
        """Get current processing statistics."""
        return self.stats
    
    async def handle_message(self, message: AgentMessage) -> AgentMessage:
        """
        Handle incoming messages from other agents.
        
        Args:
            message: Incoming agent message
            
        Returns:
            Response message
        """
        try:
            if message.message_type == MessageType.REQUEST:
                content = message.content
                
                if content.get("action") == "process_document":
                    file_path = Path(content["file_path"])
                    metadata = await self.process_document(file_path)
                    
                    return AgentMessage(
                        sender=self.agent_type,
                        recipient=message.sender,
                        message_type=MessageType.RESPONSE,
                        content={
                            "status": "success",
                            "metadata": metadata.to_dict()
                        },
                        correlation_id=message.message_id
                    )
                
                elif content.get("action") == "search_documents":
                    query = content["query"]
                    top_k = content.get("top_k")
                    results = await self.search_documents(query, top_k)
                    
                    return AgentMessage(
                        sender=self.agent_type,
                        recipient=message.sender,
                        message_type=MessageType.RESPONSE,
                        content={
                            "status": "success",
                            "results": [
                                {
                                    "content": doc.page_content,
                                    "metadata": doc.metadata
                                } for doc in results
                            ]
                        },
                        correlation_id=message.message_id
                    )
                
                elif content.get("action") == "get_stats":
                    return AgentMessage(
                        sender=self.agent_type,
                        recipient=message.sender,
                        message_type=MessageType.RESPONSE,
                        content={
                            "status": "success",
                            "stats": self.stats.to_dict()
                        },
                        correlation_id=message.message_id
                    )
            
            # Unknown action
            return AgentMessage(
                sender=self.agent_type,
                recipient=message.sender,
                message_type=MessageType.ERROR,
                content={
                    "error": f"Unknown action: {message.content.get('action', 'none')}"
                },
                correlation_id=message.message_id
            )
            
        except Exception as e:
            self.logger.error("Message handling failed", error=str(e))
            return AgentMessage(
                sender=self.agent_type,
                recipient=message.sender,
                message_type=MessageType.ERROR,
                content={
                    "error": str(e)
                },
                correlation_id=message.message_id
            )
